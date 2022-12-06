from docx import Document

import pandas as pd
import math
import seaborn



Laptop_prices = pd.read_csv("commodity 2000-2022.csv")

print(seaborn.heatmap(Laptop_prices))
doc = Document()

doc.add_heading('DATA TOOLS', 0)

a = 5
b = 3

doc.add_heading("SIMPLE MATH")
p1 = doc.add_paragraph("lets try to do a simple math, ")

p1.add_run("lets say we have the number  ")
p1.add_run(  str(a))
p1.add_run(" and the number  ")
p1.add_run( str(b) ) 

p1.add_run(" total will be " )
p1.add_run( str(a+b) )

p2 = doc.add_paragraph("TEST DATA WILL BE EQUAL TO")



doc.save("test.docx")
'''
a = Laptop_prices.head()
a.columns[2]
print(a.shape[1])

table = doc.add_table(rows = 1 ,cols = a.shape[1])

for i in range(len(a)):
    table.rows[0].cells[i].text = a.columns[i]

for a.columns in a :
    cells = table.add_row.cells
    for j in range(a.shape[1]):
        cells[j].text = a.columns[j]

doc.save("test.docx")

'''


