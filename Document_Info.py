import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DOCUMENT_DOCX:
    def __init__(self,csv_file_path ):

        self.csv_file_path = csv_file_path

        self.dataframe = pd.read_csv(self.csv_file_path)  


        self.document = Document()

    def set_up_parahraph(self):

        self.document.add_heading('DEMO FOR DAT', 0)

        self.entry = self.document.add_paragraph('This is the demonstration of how the program will work.It is currently on a test to see the possible results. ')
        self.entry.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.entry.add_run(" The areas with important informations will be bold.").bold = True
        self.entry.add_run(" Whenever a quotation is used,they will be italic like that.").italic = True
        self.entry.add_run(" Now lets go to example info gathering paragraph lets start.")

#####SIMPLE UNADJUSTABLE INFO###

        self.document.add_heading("INFORMATIONS",level = 1)
        self.second = self.document.add_paragraph("In that report,there will be 5 main spesifications for the dataframe")
        self.second.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.second.add_run("These are ; dtypes , dfinfo, shape, describe and isnull_sum")
        
        
        



    def document_columns(self):

        self.document.add_heading("DataFrame Columns",level =2)
        self.columns_paragraph = self.document.add_paragraph("In this part , we can see the datatypes for each columns of the dataframe which are :\n")
        
        for i in range (0,len(self.dataframe.columns)):

            self.columns_paragraph.add_run(f"column  {i} == {self.dataframe.columns[i]}\n")
        

    def document_dtype(self):

        self.dtypes = self.dataframe.dtypes

        self.document.add_heading("Dataframe Dtypes", level = 2)

        self.table = self.document.add_table(rows=len(self.dataframe.columns), cols=2,  style="Table Grid" ) 

        for i in range(0,len(self.dtypes)):
            self.row = self.table.rows[i]
            self.row.cells[0].text = str(self.dataframe.columns[i])
            self.row.cells[1].text = str(self.dtypes[i])

    

    def document_dfinfo(self):

        self.dfinfo = self.dataframe.info()

        self.document.add_heading("Dataframe DFINFO", level = 2)

        self.doc_df_paragraph = self.document.add_paragraph(f"{str(self.dfinfo)}")

        



        pass

    def document_shape(self):
        pass
    def document_describe(self):
        pass

    def document_isnull_sum(self):
        pass

    
    def save_document(self):

        self.document.save('demo.docx')
